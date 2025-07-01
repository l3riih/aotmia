from __future__ import annotations

"""Content parsers for supported input formats.

Each parser class converts a *bytes* or *str* representation of an input file
into **plain text** plus optional metadata that downstream pipeline steps can
consume.
"""

from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any, Dict, Optional, Tuple

# Third-party libs (guarded imports to keep optional deps optional)
try:
    from pypdf import PdfReader  # type: ignore
except ImportError:  # pragma: no cover
    PdfReader = None  # type: ignore

try:
    import docx  # python-docx  # type: ignore
except ImportError:  # pragma: no cover
    docx = None  # type: ignore

try:
    from bs4 import BeautifulSoup  # type: ignore
except ImportError:  # pragma: no cover
    BeautifulSoup = None  # type: ignore

import markdown as _md  # markdown is light dependency

# Additional imports for new parsers
try:
    import ebooklib
    from ebooklib import epub
except ImportError:  # pragma: no cover
    ebooklib = None  # type: ignore
    epub = None  # type: ignore

try:
    import httpx
except ImportError:  # pragma: no cover
    httpx = None  # type: ignore


class ParserError(Exception):
    """Raised when a parser cannot process the given input."""


class BaseParser(ABC):
    """Abstract parser class."""

    content_type: str  # MIME-like identifier or shorthand (e.g. "pdf")

    @abstractmethod
    def parse(self, data: bytes | str, *, filename: str | None = None) -> Tuple[str, Dict[str, Any]]:
        """Return `(plain_text, metadata)` given raw *data*.
        `data` may be bytes (e.g. file content) or string for already decoded
        text. Implementations should raise `ParserError` if parsing fails.
        """


class TxtParser(BaseParser):
    content_type = "text/plain"

    def parse(self, data: bytes | str, *, filename: str | None = None):
        if isinstance(data, bytes):
            try:
                text = data.decode("utf-8")
            except UnicodeDecodeError as exc:
                raise ParserError("Unable to decode text file as UTF-8") from exc
        else:
            text = data
        metadata = {"filename": filename, "length": len(text)}
        return text, metadata


class PdfParser(BaseParser):
    content_type = "application/pdf"

    def parse(self, data: bytes | str, *, filename: str | None = None):
        if PdfReader is None:
            raise ParserError("pypdf not installed")
        if isinstance(data, str):
            raise ParserError("PDF parser expects bytes content")
        try:
            from io import BytesIO
            reader = PdfReader(BytesIO(data))
            pages_text: list[str] = [page.extract_text() or "" for page in reader.pages]
            text = "\n".join(pages_text)
        except Exception as exc:  # pragma: no cover
            raise ParserError(f"Failed to parse PDF: {exc}") from exc
        metadata = {"filename": filename, "pages": len(reader.pages)}
        return text, metadata


class DocxParser(BaseParser):
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def parse(self, data: bytes | str, *, filename: str | None = None):
        if docx is None:
            raise ParserError("python-docx not installed")
        if isinstance(data, str):
            raise ParserError("DOCX parser expects bytes content")
        try:
            from io import BytesIO
            document = docx.Document(BytesIO(data))
            text = "\n".join([p.text for p in document.paragraphs])
        except Exception as exc:  # pragma: no cover
            raise ParserError(f"Failed to parse DOCX: {exc}") from exc
        metadata = {"filename": filename, "paragraphs": len(document.paragraphs)}
        return text, metadata


class HtmlParser(BaseParser):
    content_type = "text/html"

    def parse(self, data: bytes | str, *, filename: str | None = None):
        if BeautifulSoup is None:
            raise ParserError("beautifulsoup4 not installed")
        if isinstance(data, bytes):
            try:
                data = data.decode("utf-8")
            except UnicodeDecodeError as exc:
                raise ParserError("Unable to decode HTML file as UTF-8") from exc
        soup = BeautifulSoup(data, "html.parser")
        # remove script/style tags
        for tag in soup(["script", "style"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
        metadata = {"filename": filename, "title": soup.title.string if soup.title else None}
        return text, metadata


class MarkdownParser(BaseParser):
    content_type = "text/markdown"

    def parse(self, data: bytes | str, *, filename: str | None = None):
        if isinstance(data, bytes):
            try:
                data = data.decode("utf-8")
            except UnicodeDecodeError as exc:
                raise ParserError("Unable to decode Markdown file as UTF-8") from exc
        # Convert markdown to plain text by stripping formatting.
        html = _md.markdown(data)
        if BeautifulSoup is None:
            text = html  # fallback
        else:
            soup = BeautifulSoup(html, "html.parser")
            text = soup.get_text(separator="\n")
        metadata = {"filename": filename, "length": len(text)}
        return text, metadata


class EpubParser(BaseParser):
    content_type = "application/epub+zip"

    def parse(self, data: bytes | str, *, filename: str | None = None):
        if ebooklib is None or epub is None:
            raise ParserError("ebooklib not installed")
        if isinstance(data, str):
            raise ParserError("EPUB parser expects bytes content")
        
        try:
            from io import BytesIO
            import tempfile
            
            # EPUB library needs a file, so we use a temporary file
            with tempfile.NamedTemporaryFile(suffix='.epub', delete=False) as tmp_file:
                tmp_file.write(data)
                tmp_file.flush()
                
                # Read EPUB
                book = epub.read_epub(tmp_file.name)
                
                # Extract text from all chapters
                chapters_text = []
                for item in book.get_items():
                    if item.get_type() == ebooklib.ITEM_DOCUMENT:
                        content = item.get_content()
                        if BeautifulSoup:
                            soup = BeautifulSoup(content, 'html.parser')
                            text = soup.get_text(separator="\n")
                            chapters_text.append(text)
                        else:
                            chapters_text.append(content.decode('utf-8'))
                
                # Clean up temp file
                import os
                os.unlink(tmp_file.name)
                
            text = "\n\n".join(chapters_text)
            
        except Exception as exc:  # pragma: no cover
            raise ParserError(f"Failed to parse EPUB: {exc}") from exc
            
        metadata = {
            "filename": filename,
            "title": book.get_metadata('DC', 'title')[0][0] if book.get_metadata('DC', 'title') else None,
            "author": book.get_metadata('DC', 'creator')[0][0] if book.get_metadata('DC', 'creator') else None,
            "chapters": len(chapters_text)
        }
        return text, metadata


class UrlParser(BaseParser):
    content_type = "text/url"

    def parse(self, data: bytes | str, *, filename: str | None = None):
        if httpx is None:
            raise ParserError("httpx not installed")
        if BeautifulSoup is None:
            raise ParserError("beautifulsoup4 not installed")
            
        # data should be the URL string
        if isinstance(data, bytes):
            url = data.decode("utf-8")
        else:
            url = data
            
        try:
            import asyncio
            
            async def fetch_content():
                async with httpx.AsyncClient() as client:
                    response = await client.get(url, follow_redirects=True)
                    response.raise_for_status()
                    return response.text
            
            # Run async fetch
            html_content = asyncio.run(fetch_content())
            
            # Parse HTML and extract text
            soup = BeautifulSoup(html_content, "html.parser")
            
            # Remove script, style, and other non-content elements
            for tag in soup(["script", "style", "nav", "header", "footer", "aside"]):
                tag.decompose()
                
            # Try to find main content
            main_content = soup.find("main") or soup.find("article") or soup
            text = main_content.get_text(separator="\n")
            
            # Clean up text
            lines = [line.strip() for line in text.split('\n')]
            text = '\n'.join(line for line in lines if line)
            
        except Exception as exc:  # pragma: no cover
            raise ParserError(f"Failed to parse URL: {exc}") from exc
            
        metadata = {
            "url": url,
            "title": soup.title.string if soup.title else None,
            "length": len(text)
        }
        return text, metadata


SUPPORTED_PARSERS: dict[str, BaseParser] = {
    TxtParser.content_type: TxtParser(),
    PdfParser.content_type: PdfParser(),
    DocxParser.content_type: DocxParser(),
    HtmlParser.content_type: HtmlParser(),
    MarkdownParser.content_type: MarkdownParser(),
    EpubParser.content_type: EpubParser(),
    UrlParser.content_type: UrlParser(),
}


def detect_content_type(filename: str | None, content_type_header: str | None) -> str:
    """Best-effort detection of content type based on filename or header."""
    if content_type_header:
        return content_type_header.split(";")[0].strip()
    if not filename:
        return "text/plain"
    
    # Check if it's a URL
    if filename.startswith(('http://', 'https://')):
        return UrlParser.content_type
        
    ext = Path(filename).suffix.lower()
    match ext:
        case ".pdf":
            return PdfParser.content_type
        case ".docx":
            return DocxParser.content_type
        case ".html" | ".htm":
            return HtmlParser.content_type
        case ".md":
            return MarkdownParser.content_type
        case ".epub":
            return EpubParser.content_type
        case _:
            return "text/plain"


def parse_content(data: bytes | str, *, filename: str | None = None, content_type: str | None = None) -> tuple[str, Dict[str, Any]]:
    """Convenience helper to parse arbitrary input using appropriate parser."""
    ctype = detect_content_type(filename, content_type)
    parser = SUPPORTED_PARSERS.get(ctype)
    if not parser:
        raise ParserError(f"Unsupported content type: {ctype}")
    return parser.parse(data, filename=filename) 